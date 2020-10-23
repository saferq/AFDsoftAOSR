# -*- coding: utf-8 -*-
import datetime
import pprint as pp
import docx as dc
import re

from res import afd_sqlite as db


class Filling_act():
    """ Создание АОСР из словарей данных """
    def __init__(self):
        """ Конструктор """
        print("Filling_act")
        self.dtime = datetime.datetime.today()
        self.db = db.Work_with_datebase()

    def creat_aosr(self, path, act_number):
        """ Создание акта скрытых работ """
        print("Filling_act.creat_aosr2")
     # Методы для работы с word
        self.document = dc.Document()
        self.cm = dc.shared.Cm
        self.pt = dc.shared.Pt
        self.wd_align = dc.enum.text.WD_PARAGRAPH_ALIGNMENT
        self.wd_break = dc.enum.text.WD_BREAK
        self.wd_style = dc.enum.style.WD_STYLE_TYPE
        self.wd_align = dc.enum.text.WD_PARAGRAPH_ALIGNMENT
        self.wd_align_v = dc.enum.table.WD_ALIGN_VERTICAL
        self.wd_tb_align = dc.enum.table.WD_TABLE_ALIGNMENT
        self.styles = self.document.styles
     # Поля листа
        self.section = self.document.sections[0]
        self.section.bottom_margin = self.cm(1)
        self.section.top_margin = self.cm(1)
        self.section.left_margin = self.cm(1.5)
        self.section.right_margin = self.cm(1.5)
     # Стили
      # Стиль текста style='style_zag'
        self.style_txt = self.styles.add_style(
            'style_zag', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.JUSTIFY
        p_format.line_spacing = 1
        p_format.keep_with_next = True
        font = self.style_txt.font
        font.name = 'Franklin Gothic Book'
        font.size = self.pt(12)
        font.underline = False
      # Стиль текста style='style_type'
        self.style_txt = self.styles.add_style(
            'style_type', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(5)
        p_format.space_after = self.pt(5)
        p_format.alignment = self.wd_align.CENTER
        p_format.line_spacing = 1
        p_format.keep_with_next = True
        font = self.style_txt.font
        font.name = 'Franklin Gothic Book'
        font.size = self.pt(12)
        font.underline = False
        font.bold = True
      # Стиль текста style='style_txt'
        self.style_txt = self.styles.add_style(
            'style_txt', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.JUSTIFY
        p_format.line_spacing = 1
        p_format.keep_with_next = True
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Franklin Gothic Book'
        font.size = self.pt(12)
        font.italic = True
        font.underline = True
      # Стиль текста style='style_ops' для описания в скорбках
        self.style_txt = self.styles.add_style(
            'style_ops', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.CENTER
        p_format.line_spacing = 1
        p_format.keep_with_next = False
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Franklin Gothic Book'
        font.size = self.pt(6)
        font.italic = True
        font.underline = False
      # Стиль текста style='style_ops_head' для шапки
        self.style_txt = self.styles.add_style(
            'style_ops_head', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.RIGHT
        p_format.line_spacing = 1
        p_format.keep_with_next = False
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Franklin Gothic Book'
        font.size = self.pt(6)
        font.italic = True
        font.underline = False
      # Стиль текста style = 'style_table'
        self.style_txt = self.styles.add_style(
            'style_table', self.wd_style.TABLE)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.CENTER
        p_format.line_spacing = 1
        p_format.keep_with_next = False
        # p_tab = self.style_txt.paragraph_format.tab_stops
        # p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Franklin Gothic Book'
        font.size = self.pt(12)
        font.italic = True
        font.underline = True
      # Стиль текста style = 'style_table_vsn'
        self.style_txt = self.styles.add_style(
            'style_table_vsn', self.wd_style.TABLE)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.CENTER
        p_format.line_spacing = 1
        p_format.keep_with_next = False
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Franklin Gothic Book'
        font.size = self.pt(10)
        font.italic = True
        font.underline = True
     # Титульный лист
      # Шапка
        text = """ РД–11–02–2006   Приложение № 3 \nк Требованиям к составу и порядку ведения исполнительной документации при строительстве,   реконструкции, капитальном ремонте объектов капитального строительства и требования, предъявляемые к актам освидетельствования работ, конструкций, участков сетей инженерно-технического обеспечения, утвержденным Приказом Федеральной службы по экологическому, технологическому и атомному надзору от 26 декабря 2006 г. N 1128 """,
        self.p = self.document.add_paragraph(text, style='style_ops_head')
      # Объект
        text_t = 'Объект капитального строительства:'
        text_d = '(наименование, почтовый или строительный адрес объекта капитального строительства)'
        query = "SELECT object FROM Объект WHERE rowid==3;"
        text = self.db.db_query_fetchone(query)
        self.insert_1(text_t, text_d, text, '\t ')
      # Застройщик 
        text_t = 'Застройщик (технический заказчик, эксплуатирующая организация или региональный оператор):'
        text_d = '(фамилия, имя, отчество , адрес места жительства, ОГРНИП, ИНН индивидуального предпринимателя, наименование, ОГРН, ИНН, место нахождения, юридического лица, телефон/факс, наименование, ОГРН, ИНН саморегулируемой организации, членом которой является  – для индивидуальных предпринимателей и юридических лиц; фамилия, имя, отчество , паспортные данные, адрес места жительства, телефон/факс – для физических лиц, не являющихся индивидуальными предпринимателями)'
        query = "SELECT custumer FROM Объект WHERE rowid==3;"
        text = self.db.db_query_fetchone(query)
        self.insert_1(text_t, text_d, text, '\t ')
      # Генподрядчик
        text_t = 'Лицо, осуществляющее строительство:'
        text_d = '(фамилия, имя, отчество , адрес места жительства, ОГРНИП, ИНН индивидуального предпринимателя, наименование, ОГРН, ИНН, место нахождения, юридического лица, телефон/факс, наименование, ОГРН, ИНН саморегулируемой организации, членом которой является )'
        query = "SELECT builder FROM Объект WHERE rowid==3;"
        text = self.db.db_query_fetchone(query)
        self.insert_1(text_t, text_d, text, '\t ')
      # Проектировщик
        text_t = 'Лицо, осуществляющее подготовку проектной документации:'
        text_d = '(фамилия, имя, отчество , адрес места жительства, ОГРНИП, ИНН индивидуального предпринимателя, наименование, ОГРН, ИНН, место нахождения, юридического лица, телефон/факс, наименование, ОГРН, ИНН саморегулируемой организации, членом которой является )'
        query = "SELECT project FROM Объект WHERE rowid==3;"
        text = self.db.db_query_fetchone(query)
        self.insert_1(text_t, text_d, text, '\t ')
     # Номер дата акта (таблица)
        self.p = self.document.add_paragraph(style='style_type')
        self.run = self.p.add_run("АКТ\nосвидетельствования скрытых работ")
        # SQL
        query = f"SELECT number FROM АОСР WHERE number=='{act_number}' LIMIT 1;"
        text1 = self.get_one_search(query)
        query = f"SELECT date1 FROM АОСР WHERE number=='{act_number}' LIMIT 1;"
        text2 = self.get_one_search(query)
        # Таблица
        table = self.document.add_table(rows=2, cols=3)
        table.style = 'style_table'
        table.alignment = self.wd_tb_align.CENTER
        cell = table.cell(0, 0)
        cell.text = text1 # Номер
        cell = table.cell(0, 2)
        cell.text = text2 + ' г.' # Дата
        cell = table.cell(1, 2)
        cell.text = '(дата составления акта)'
        cell.paragraphs[0].paragraph_format.alignment = self.wd_align.CENTER
        cell.paragraphs[0].runs[0].font.size = self.pt(6)
        cell.paragraphs[0].runs[0].font.italic = False
        cell.paragraphs[0].runs[0].font.underline = False

     # Представители
      # ZK Представитель застройщика
        text_t = 'Представитель застройщика (технического заказчика, эксплуатирующей организации или регионального оператора) по вопросам строительного контроля'
        text_d = '(должность, фамилия, инициалы, идентификационный номер в национальном реестре специалистов в области строительства , реквизиты распорядительного документа, подтверждающие полномочия, с указанием наименования, ОГРН, ИНН, места нахождения юридического лица ,фамилии, имени, отчества , адреса места жительства, ОГРНИП, ИНН индивидуального предпринимателя )'
        text = self.get_fio('fio_full', 'ZK', act_number)
        self.insert_1(text_t, text_d, text, '')
      # GP Представитель лица, осуществляющего строительство
        text_t = 'Представитель лица, осуществляющего строительство:'
        text_d = '(должность, фамилия, инициалы, реквизиты распорядительного документа, подтверждающие полномочия)'
        text = self.get_fio('fio_full', 'GP', act_number)
        self.insert_1(text_t, text_d, text, '')
      # SKK Представитель лица, осуществляющего строительный контроль
        text_t = 'Представитель лица, осуществляющего строительство, по вопросам строительного контроля (специалист по организации строительства)'
        text_d = '(должность, фамилия, инициалы, идентификационный номер в национальном реестре специалистов в области строительства, реквизиты распорядительного документа, подтверждающие полномочия)'
        text = self.get_fio('fio_full', 'SKK', act_number)
        self.insert_1(text_t, text_d, text, '')
      # AN Представитель лица, осуществляющего подготовку проектной документации
        text_t = 'Представитель лица, осуществляющего подготовку проектной документации'
        text_d = '(должность, фамилия, инициалы, идентификационный номер в национальном реестре специалистов в области строительства, реквизиты распорядительного документа, подтверждающие полномочия)'
        text = self.get_fio('fio_full', 'AN', act_number)
        self.insert_1(text_t, text_d, text, '')
      # PD Представитель лица, выполнившего работы, подлежащие освидетельствованию
        text_t = 'Представитель лица, выполнившего работы, подлежащие освидетельствованию'
        text_d = '(должность, фамилия, инициалы, реквизиты распорядительного документа, подтверждающие полномочия, с указанием наименования, ОГРН, ИНН, места нахождения юридического лица, фамилии, имени, отчества , адреса места жительства, ОГРНИП, ИНН индивидуального предпринимателя)'
        text = self.get_fio('fio_full', 'PD', act_number)
        self.insert_1(text_t, text_d, text, '')
      # SK Иные представители лиц, технадзоры
        text_t = 'а также иные представители лиц, участвующих в освидетельствовании:'
        text_d = '(должность с указанием наименования организации, фамилия, инициалы, реквизиты распорядительного документа, подтверждающие полномочия)'
        text = self.get_fio('fio_full', 'SK', act_number)
        self.insert_1(text_t, text_d, text, '')
 
     # Произвели осмотр работ
        query = "SELECT build_name FROM Объект LIMIT 1"
        text = self.db.db_query(query)[0][0]
        self.p = self.document.add_paragraph(style='style_txt')
        self.run = self.p.add_run(f"произвели осмотр работ, выполненных {text}\t ")
        self.p = self.document.add_paragraph('(наименование лица, выполнившего работы, подлежащие освидетельствованию)', style='style_ops')
        self.p = self.document.add_paragraph('и составили настоящий акт о нижеследующем:', style='style_zag')
        self.p.paragraph_format.keep_with_next = False
     # 1. К освидетельствованию предъявлены
        text_t = '1. К освидетельствованию предъявлены следующие работы:'
        text_d = '(наименование скрытых работ)'
        query = f"SELECT name FROM АОСР WHERE number=='{act_number}' AND type=='АОСР';"
        table = self.db.db_query(query)
        text = ''
        for row in table:
            for i in row:
                text += i 
                text += "\t " 
        self.insert_1(text_t, text_d, text, '')
        self.p.paragraph_format.keep_with_next = True
     # 2. Работы выполнены по проектной документации:
        text_t = '2. Работы выполнены по проектной документации:'
        text_d = '(номер, другие реквизиты чертежа, наименование проектной и/или рабочей документации, сведения о лицах, осуществляющих подготовку раздела проектной и/или рабочей документации)'
        query = f"SELECT proj FROM АОСР WHERE number=='{act_number}' AND type=='АОСР' LIMIT 1;"
        table = self.db.db_query(query)
        text = table[0][0]
        self.insert_1(text_t, text_d, text, '')
        self.p.paragraph_format.keep_with_next = True
     # 3. При выполнении работ применены:
        text_t = '3. При выполнении работ применены:'
        text_d = '(наименование строительных материалов (изделий), реквизиты сертификатов и/или других документов, подтверждающих качество и безопасность)'
        query = f"SELECT name FROM АОСР WHERE number=='{act_number}' AND type=='ВК';"
        table = self.db.db_query(query)
        text = ''
        for row in table:
            for i in row:
                text += i 
                text += "; " 
        text += "\t " 
        self.insert_1(text_t, text_d, text, '')
     # 4. Предъявлены документы
        text_t = '4. Предъявлены документы, подтверждающие соответствие работ предъявляемым к ним требованиям:'
        text_d = '(исполнительные схемы и чертежи, результаты экспертиз, обследований, лабораторных и иных испытаний выполненных работ, проведенных в процессе строительного контроля)'
        query = f"SELECT proj FROM АОСР WHERE number=='{act_number}' AND type=='ВК' LIMIT 1;"
        table = self.db.db_query(query)
        text1 = table[0][0]
        query = f"SELECT pril2 FROM АОСР WHERE number=='{act_number}' AND type=='ВК' LIMIT 1;"
        table = self.db.db_query(query)
        text2 = table[0][0]
        text = text1 + ", " + text2 + "\t "
        self.insert_1(text_t, text_d, text, '')
     # 5.Даты:
        query = f"SELECT date2 FROM АОСР WHERE number=='{act_number}' AND type=='АОСР' LIMIT 1;"
        table = self.db.db_query(query)
        date_s = table[0][0]
        query = f"SELECT date3 FROM АОСР WHERE number=='{act_number}' AND type=='АОСР' LIMIT 1;"
        table = self.db.db_query(query)
        date_f = table[0][0]
      # Начало работ
        self.p = self.document.add_paragraph(style='style_zag')
        self.tabs = self.p.paragraph_format.tab_stops
        self.tabs.add_tab_stop(self.cm(2))
        self.tabs.add_tab_stop(self.cm(6))
        self.run = self.p.add_run(f'5.Даты:\tначала работ:\t')
        self.run = self.p.add_run(date_s)
        self.run.font.italic = True
        self.run.font.underline = True
        self.run = self.p.add_run(' г.')
      # Окончание работ
        self.p = self.document.add_paragraph(style='style_zag')
        self.tabs = self.p.paragraph_format.tab_stops
        self.tabs.add_tab_stop(self.cm(2))
        self.tabs.add_tab_stop(self.cm(6))
        self.run = self.p.add_run(f'\tокончания работ:\t')
        self.run = self.p.add_run(date_f)
        self.run.font.italic = True
        self.run.font.underline = True
        self.run = self.p.add_run(' г.')
     # 6. Работы выполнены в соответствии с проектной документацией
        text_t = '6. Работы выполнены в соответствии с проектной документацией '
        text_d = '(наименования и структурные единицы технических регламентов, иных нормативных правовых актов, разделы проектной и/или рабочей документации)'
        query = f"SELECT pril2 FROM АОСР WHERE number=='{act_number}' AND type=='АОСР' LIMIT 1;"
        table = self.db.db_query(query)
        text = table[0][0]
        self.insert_1(text_t, text_d, text, '\t ')
     # 7. Разрешается производство последующих работ'
        text_t = '7. Разрешается производство последующих работ'
        text_d = '(наименование работ, конструкций, участков сетей инженерно-технического обеспечения)'
        query = f"SELECT razr FROM АОСР WHERE number=='{act_number}' AND type=='АОСР' LIMIT 1;"
        table = self.db.db_query(query)
        text = table[0][0]
        self.insert_1(text_t, text_d, text, '\t ')
     # Дополнительные сведения
        self.p = self.document.add_paragraph('Дополнительные сведения ', style='style_zag')
        self.p = self.document.add_paragraph(style='style_txt')
        self.run = self.p.add_run(f"Отсутствуют \t ")
        self.p = self.document.add_paragraph('Акт составлен в 3 экземплярах. ', style='style_zag')
     # Приложения:
        text_t = 'Приложения:'
        text_d = '(исполнительные схемы и чертежи, результаты экспертиз, обследований, лабораторных и иных испытаний)'
        query = f"SELECT proj FROM АОСР WHERE number=='{act_number}' AND type=='ВК' LIMIT 1;"
        table = self.db.db_query(query)
        text1 = table[0][0]
        query = f"SELECT pril2 FROM АОСР WHERE number=='{act_number}' AND type=='ВК' LIMIT 1;"
        table = self.db.db_query(query)
        text2 = table[0][0]
        text = text1 + ", " + text2 + "\t "
        self.insert_1(text_t, text_d, text, '')

     # Подписанты
      # ZK Представитель застройщика
        text_t = 'Представитель застройщика (технического заказчика, эксплуатирующей организации или регионального оператора) по вопросам строительного контроля '
        text_d = '(фамилия, инициалы, подпись)'
        text = self.get_fio('fio', 'ZK', act_number)
        self.insert_1(text_t, text_d, text, '')
      # GP Представитель лица, осуществляющего строительство
        text_t = 'Представитель лица, осуществляющего строительство'
        text_d = '(фамилия, инициалы, подпись)'
        text = self.get_fio('fio', 'GP', act_number)
        self.insert_1(text_t, text_d, text, '')
      # SKK Представитель лица, осуществляющего строительный контроль
        text_t = 'Представитель лица, осуществляющего строительство, по вопросам строительного контроля (специалист по организации строительства)'
        text_d = '(фамилия, инициалы, подпись)'
        text = self.get_fio('fio', 'SKK', act_number)
        self.insert_1(text_t, text_d, text, '')
      # AN Представитель лица, осуществляющего подготовку проектной документации
        text_t = 'Представитель лица, осуществляющего подготовку проектной документации'
        text_d = '(фамилия, инициалы, подпись)'
        text = self.get_fio('fio', 'AN', act_number)
        self.insert_1(text_t, text_d, text, '')
      # PD Представитель лица, выполнившего работы, подлежащие освидетельствованию
        text_t = 'Представитель лица, выполнившего работы, подлежащие освидетельствованию,'
        text_d = '(фамилия, инициалы, подпись)'
        text = self.get_fio('fio', 'PD', act_number)
        self.insert_1(text_t, text_d, text, '')
      # SK Представитель иных лиц:
        text_t = 'Представитель иных лиц:'
        text_d = '(фамилия, инициалы, подпись)'
        text = self.get_fio('fio', 'SK', act_number)
        self.insert_1(text_t, text_d, text, '')

     # Сохранение документа
        # txt_date = self.dtime.strftime('%m%d-%H-%M-%S')
        txt_date = self.dtime.strftime('%m%d')
        query = f"SELECT name FROM АОСР WHERE number=='{act_number}' AND type=='АОСР';"
        table = self.db.db_query(query)
        txt_name = table[0][0]
        txt_name = re.compile('[^a-zA-Zа-яА-Я0-9 ]').sub("", txt_name)
        txt_name = txt_name.replace(' ', '_').replace('__', '_')
        print(f"Сохранение АОСР: AOSR_{act_number}_{txt_name}.docx")
        self.document.save(path + f"АОСР_{act_number}_{txt_name}.docx")

  ############# ВСН 012-88 форма № 3.3 #############
    # def creat_vk_vsn88 (self, vk_title, vk_general, vk_fio):
    def creat_vk_vsn88(self, title, general, fio, path):
        """ 
        Создание акта входного контроля\n
        по ВСН 012-88 форма № 3.3 
        """
        print("Filling_act.creat_vk_vsn88")
     # Методы для работы с word
        self.document = dc.Document()
        self.cm = dc.shared.Cm
        self.pt = dc.shared.Pt
        self.wd_align = dc.enum.text.WD_PARAGRAPH_ALIGNMENT
        self.wd_break = dc.enum.text.WD_BREAK
        self.wd_style = dc.enum.style.WD_STYLE_TYPE
        self.wd_align = dc.enum.text.WD_PARAGRAPH_ALIGNMENT
        self.wd_align_v = dc.enum.table.WD_ALIGN_VERTICAL
        self.wd_tb_align = dc.enum.table.WD_TABLE_ALIGNMENT
        self.styles = self.document.styles
     # Поля листа
        self.section = self.document.sections[0]
        self.section.bottom_margin = self.cm(1)
        self.section.top_margin = self.cm(1)
        self.section.left_margin = self.cm(1.5)
        self.section.right_margin = self.cm(1.5)
     # Стили
      # Стиль текста style='style_zag'
        self.style_txt = self.styles.add_style(
            'style_zag', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.JUSTIFY
        p_format.line_spacing = 1
        p_format.keep_with_next = True
        font = self.style_txt.font
        font.name = 'Times New Roman'
        font.size = self.pt(12)
        font.underline = False
      # Стиль текста style='style_type'
        self.style_txt = self.styles.add_style(
            'style_type', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(5)
        p_format.space_after = self.pt(5)
        p_format.alignment = self.wd_align.CENTER
        p_format.line_spacing = 1
        p_format.keep_with_next = True
        font = self.style_txt.font
        font.name = 'Times New Roman'
        font.size = self.pt(12)
        font.underline = False
        font.bold = True
      # Стиль текста style='style_txt'
        self.style_txt = self.styles.add_style(
            'style_txt', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.JUSTIFY
        p_format.line_spacing = 1
        p_format.keep_with_next = True
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Times New Roman'
        font.size = self.pt(12)
        font.italic = True
        font.underline = True
      # Стиль текста style='style_ops' для описания в скорбках
        self.style_txt = self.styles.add_style(
            'style_ops', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.CENTER
        p_format.line_spacing = 1
        p_format.keep_with_next = False
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Times New Roman'
        font.size = self.pt(7)
        font.italic = True
        font.underline = False
      # Стиль текста style='style_ops_head' для шапки
        self.style_txt = self.styles.add_style(
            'style_ops_head', self.wd_style.PARAGRAPH)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.RIGHT
        p_format.line_spacing = 1
        p_format.keep_with_next = False
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Times New Roman'
        font.size = self.pt(6)
        font.italic = True
        font.underline = False
      # Стиль текста style = 'style_table'
        self.style_txt = self.styles.add_style(
            'style_table', self.wd_style.TABLE)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.CENTER
        p_format.line_spacing = 1
        p_format.keep_with_next = False
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Times New Roman'
        font.size = self.pt(12)
        font.italic = True
        font.underline = True
      # Стиль текста style = 'style_table_vsn'
        self.style_txt = self.styles.add_style(
            'style_table_vsn', self.wd_style.TABLE)
        p_format = self.style_txt.paragraph_format
        p_format.space_before = self.pt(0)
        p_format.space_after = self.pt(0)
        p_format.alignment = self.wd_align.CENTER
        p_format.line_spacing = 1
        p_format.keep_with_next = False
        p_tab = self.style_txt.paragraph_format.tab_stops
        p_tab.add_tab_stop(self.cm(18.5))
        font = self.style_txt.font
        font.name = 'Times New Roman'
        font.size = self.pt(10)
        font.italic = True
        font.underline = True

     # Шапка ВСН 012-88 (таблица)
        table = self.document.add_table(rows=8, cols=25)
        table.style = 'style_table_vsn'
        table.alignment = self.wd_tb_align.CENTER
     # строка 1 форма
      # Колонка 2
        cell = table.cell(0, 12)
        cell.merge(table.cell(0, 24))
        cell.vertical_alignment = self.wd_align_v.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = self.wd_align.RIGHT
        run = p.add_run('Форма N 3.3 (рекомендуемая) ')
        run.font.size = self.pt(10)
        run.font.underline = False
        run.font.italic = False
     # строка 2 Министерство / Основание
      # Колонка 1
        cell = table.cell(1, 0)
        cell.merge(table.cell(1, 11))
        cell.vertical_alignment = self.wd_align_v.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = self.wd_align.LEFT
        run = p.add_run('Министерство: ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(title['ministry'][0])
        run.font.underline = True
        run.font.italic = True
        # 3
        cell = table.cell(1, 12)
        cell.merge(table.cell(2, 14))
        cell.text = 'Основание: '
        cell.paragraphs[0].paragraph_format.alignment = self.wd_align.LEFT
        cell.paragraphs[0].runs[0].font.underline = False
        cell.paragraphs[0].runs[0].font.italic = False
        cell.vertical_alignment = self.wd_align_v.CENTER
        # 4
        cell = table.cell(1, 15)
        cell.merge(table.cell(1, 24))
        cell.text = 'ВСН 012-88 (Часть II) '
        cell.paragraphs[0].paragraph_format.alignment = self.wd_align.LEFT
        cell.paragraphs[0].runs[0].font.underline = True
        cell.paragraphs[0].runs[0].font.italic = False
        cell.vertical_alignment = self.wd_align_v.BOTTOM
     # строка 3 Объединение / Основание
        # 1
        cell = table.cell(2, 0)
        cell.merge(table.cell(3, 11))
        cell.vertical_alignment = self.wd_align_v.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = self.wd_align.LEFT
        run = p.add_run('Объединение, трест ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(title['union'][0])
        run.font.underline = True
        run.font.italic = True
        # 4
        cell = table.cell(2, 15)
        cell.merge(table.cell(2, 24))
        cell.text = 'Миннефтегазстрой '
        cell.paragraphs[0].paragraph_format.alignment = self.wd_align.LEFT
        cell.paragraphs[0].runs[0].font.underline = False
        cell.paragraphs[0].runs[0].font.italic = False
        cell.vertical_alignment = self.wd_align_v.TOP
     # Cтрока 5 СМУ, СУ, ПМК, КТП / Строительство
      # Колонка 1
        cell = table.cell(4, 0)
        cell.merge(table.cell(5, 11))
        cell.vertical_alignment = self.wd_align_v.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = self.wd_align.LEFT
        run = p.add_run('СМУ, СУ, ПМК, КТП: ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(title['contractor2'][0])
        run.font.underline = True
        run.font.italic = True
      # Колонка 2
        cell = table.cell(4, 12)
        cell.merge(table.cell(5, 24))
        cell.vertical_alignment = self.wd_align_v.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = self.wd_align.LEFT
        run = p.add_run('Строительство: ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(title['construction'][0])
        run.font.underline = True
        run.font.italic = True
     # Cтрока 6 Участок / Объект
      # Колонка 1
        cell = table.cell(6, 0)
        cell.merge(table.cell(7, 11))
        cell.vertical_alignment = self.wd_align_v.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = self.wd_align.LEFT
        run = p.add_run('Участок: ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(title['plot'][0])
        run.font.underline = True
        run.font.italic = True
        # Колонка 2
        cell = table.cell(6, 12)
        cell.merge(table.cell(7, 24))
        cell.vertical_alignment = self.wd_align_v.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = self.wd_align.LEFT
        run = p.add_run('Объект: ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(title['object'][0])
        run.font.underline = True
        run.font.italic = True
     # Основная часть
     # Номер дата материал
        def fill_number_date(number, date, name1, name2, name3):
            p = self.document.add_paragraph(style='style_zag')
            p.paragraph_format.alignment = self.wd_align.CENTER
            run = p.add_run(f"АКТ N ")
            run.font.bold = True
            run = p.add_run(number)
            run.font.underline = True
            run.font.italic = True
            run = p.add_run(f" о результатах проверки изделий")
            run.font.bold = True
            p = self.document.add_paragraph(style='style_txt')
            run = p.add_run(f'{name1} - {name2}{name3}\t ',)
            p = self.document.add_paragraph(style='style_ops')
            run = p.add_run(f'(вид изделий)',)
            p = self.document.add_paragraph(style='style_zag')
            p.paragraph_format.alignment = self.wd_align.CENTER
            run = p.add_run(f"на соответствие техдокументации от ")
            run.font.bold = True
            run = p.add_run(date)
            run.font.underline = True
            run.font.italic = True
        fill_number_date(
            general['number_vk'],
            general['date_vk'],
            general['name1_vk'],
            general['name2_vk'],
            general['name3_vk']
        )
     # Ответственные лица
        def fill_fio_title(name1, fio_d, id):
            """ 
            name - Описание ответственного лица \n
            fio - Словарь ответственных лиц \n
            id - Идентификатор ответственного лица \n
             """
            p = self.document.add_paragraph(style='style_zag')
            run = p.add_run(f'{name1}')
            p = self.document.add_paragraph(style='style_txt')
            iter = fio_d[id][0]
            for n in iter[4:8]:
                run = p.add_run(f'{n} ')
            run = p.add_run(f'\t ')
            p = self.document.add_paragraph(style='style_ops')
            run = p.add_run(f'(должность, организация, фамилия, инициалы)',)
        # GP
        fill_fio_title(
            'Составлен представителями: строительной организации', fio, 'GP')
        # SKK
        fill_fio_title('службы контроля качества', fio, 'SKK')
        # ZK
        fill_fio_title('заказчика', fio, 'ZK')
     # Контроль
        def fill_contrl(contrl, name1, name2, name3):
            p = self.document.add_paragraph(style='style_txt')
            run = p.add_run(f'в том, что произведен ')
            run.font.underline = False
            run.font.italic = False
            run = p.add_run(f'{contrl}')
            run = p.add_run(f' осмотр ')
            run.font.underline = False
            run.font.italic = False
            run = p.add_run(f'{name1} - {name2}{name3}\t ')
            p = self.document.add_paragraph(style='style_ops')
            p.paragraph_format.alignment = self.wd_align.LEFT
            tabs = p.paragraph_format.tab_stops
            tabs.add_tab_stop(self.cm(3.75))
            tabs.add_tab_stop(self.cm(11))
            run = p.add_run(f'\t(сплошной, выборочный)',)
            run = p.add_run(f'\t(наименование изделий)',)
        fill_contrl(
            general['contrl'],
            general['name1_vk'],
            general['name2_vk'],
            general['name3_vk']
        )

        def fill_one(pre, text, ops, tab):
            p = self.document.add_paragraph(style='style_txt')
            run = p.add_run(pre)
            run.font.underline = False
            run.font.italic = False
            run = p.add_run(f'{text}\t ')
            p = self.document.add_paragraph(style='style_ops')
            tabs = p.paragraph_format.tab_stops
            tabs.add_tab_stop(tab)
            run = p.add_run(f'\t{ops}',)
     # Проект
        fill_one(
            'предназначенных проектом ',
            general['pro1_vk'],
            '(номер проекта, чертежа, дата)',
            self.cm(5))
     # Участок
        fill_one(
            'для строительства на участке трубопровода ',
            general['uchs_vk'],
            '(привязка, км/ПК)',
            self.cm(8))
     # 1 Осмотр
        fill_one(
            '1. Осмотром геометрических размеров и маркировки ',
            general['name1_vk'],
            '(труб, деталей, силовых поясов и т.д.)',
            self.cm(9))
     # Паспорта сертификаты
        fill_one(
            'совместно с проектом и сопроводительной документацией ',
            general['pasp_vk'],
            '(сертификатами, паспортами)',
            self.cm(11))
     # Установлено, что
        fill_one(
            'на изделия установлено, что ',
            general['name1_vk'] + ' - ' +
            general['name2_vk'] + general['name3_vk'],
            '(трубы, детали, силовые пояса и т.д.)',
            self.cm(6))
     # по своим геометрическим размерам
        fill_one(
            'по своим геометрическим размерам ',
            general['parm_vk'],
            '(для труб указать диаметр, толщину стенки, мм, для отводов - угол изгиба, град. и т.д.)',
            self.cm(5))
     # ТУ, ГОСТ
        fill_one(
            'и номеру технических условий, указанному на изделии, ',
            'соответствует',
            '(соответствуют, не соответствуют)',
            self.cm(10))
     # Проект
        p = self.document.add_paragraph(style='style_txt')
        run = p.add_run(f'проекту, рабочие чертежи N ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(f"{general['pro1_vk']}\t ")
     # 2.Сопроводительная документация
        p = self.document.add_paragraph(style='style_txt')
        run = p.add_run(f'2.Сопроводительная документация ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(f"{general['pasp_vk']}\t ")
        p = self.document.add_paragraph(style='style_ops')
        run = p.add_run(f'\t(паспорта, сертификаты)',)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(self.cm(6))
        p = self.document.add_paragraph(style='style_txt')
        run = p.add_run(f'имеется в полном комплекте.')
        run.font.underline = False
        run.font.italic = False
     # 3. Характеристики механических свойств
        fill_one(
            '3. Характеристики механических свойств ',
            ' ',
            '(по данным сопроводительной документации, при необходимости – результатам испытаний)',
            self.cm(8))
     # Соответствуют требованиям проекта
        p = self.document.add_paragraph(style='style_txt')
        run = p.add_run(f'соответствуют требованиям проекта ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(f"{general['pro1_vk']}\t ")
     # Технических условий
        p = self.document.add_paragraph(style='style_txt')
        run = p.add_run(f'технических условий ')
        run.font.underline = False
        run.font.italic = False
        run = p.add_run(f"{general['gost_vk']}\t ")
     # Подписанты

        def fill_podpis1(name1, fio, id, general):
            """Подписатель\n
            name1 - Представитель\n
            n2 - Фамилия И.О.\n
            n3 - Дата\n
             """
            # Кто описание
            p = self.document.add_paragraph(style='style_zag')
            run = p.add_run(name1)
            # Фамилия с инициалами
            p = self.document.add_paragraph(style='style_txt')
            run = p.add_run(f'\t',)
            iter = fio[id][0][6:8]
            for n in iter:
                run = p.add_run(f'{n} ',)
            run = p.add_run(f'\t ',)
            # Пустое пространство
            run = p.add_run(f'\t ',)
            run.font.underline = False
            run.font.italic = False
            # Место для подписи
            run = p.add_run(f'\t \t',)
            # Пустое пространство
            run = p.add_run(f'\t ',)
            run.font.underline = False
            run.font.italic = False
            # Дата
            run = p.add_run(f"\t{general['date_vk']} \t ",)
            # Растасовка табов
            tabs = p.paragraph_format.tab_stops
            tabs.add_tab_stop(self.cm(4.5)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(9)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(10)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(12)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(14)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(15)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(
                self.cm(16.75)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(
                self.cm(18.5)).alignment = self.wd_tb_align.CENTER

            p = self.document.add_paragraph(style='style_ops')
            run = p.add_run(f'\t(фамилия, инициалы)\t',)
            run = p.add_run(f'\t',)
            run.font.underline = False
            run.font.italic = False
            run = p.add_run(f'\t(подпись)\t',)
            run = p.add_run(f'\t',)
            run.font.underline = False
            run.font.italic = False
            run = p.add_run(f'\t(дата)\t',)
            tabs = p.paragraph_format.tab_stops
            tabs.add_tab_stop(self.cm(4.5)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(9)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(10)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(12)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(14)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(15)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(
                self.cm(16.75)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(
                self.cm(18.5)).alignment = self.wd_tb_align.CENTER
        fill_podpis1('Представитель строительной организации ',
                     fio, 'GP', general)
        fill_podpis1('Представитель службы контроля качества ',
                     fio, 'SKK', general)
        fill_podpis1('Представитель заказчика ', fio, 'ZK', general)

        def fill_podpis2(fio, id, general):
            """ Подписатель СКК
            n2 - Должность Фамилия И.О.\n
            n3 - Дата\n
             """
            p = self.document.add_paragraph(style='style_txt')
            p.paragraph_format.alignment = self.wd_align.LEFT
            for n in fio[id]:
                run = p.add_run(f'{n[4]} {n[5]} \n',)
                run.font.underline = False
                run = p.add_run(f'{n[6]} {n[7]} \t ',)
                run = p.add_run(f'\t ',)
                run.font.underline = False
                run.font.italic = False
                run = p.add_run(f'\t \t',)
                run = p.add_run(f'\t ',)
                run.font.underline = False
                run.font.italic = False
                run = p.add_run(f"\t{general['date_vk']} \t ",)
                tabs = p.paragraph_format.tab_stops
                tabs.add_tab_stop(
                    self.cm(9)).alignment = self.wd_tb_align.CENTER
                tabs.add_tab_stop(
                    self.cm(10)).alignment = self.wd_tb_align.CENTER
                tabs.add_tab_stop(
                    self.cm(12)).alignment = self.wd_tb_align.CENTER
                tabs.add_tab_stop(
                    self.cm(14)).alignment = self.wd_tb_align.CENTER
                tabs.add_tab_stop(
                    self.cm(15)).alignment = self.wd_tb_align.CENTER
                tabs.add_tab_stop(
                    self.cm(16.75)).alignment = self.wd_tb_align.CENTER
                tabs.add_tab_stop(
                    self.cm(18.5)).alignment = self.wd_tb_align.CENTER

            p = self.document.add_paragraph(style='style_ops')
            run = p.add_run(f'\t(фамилия, инициалы)\t',)
            run = p.add_run(f'\t',)
            run.font.underline = False
            run.font.italic = False
            run = p.add_run(f'\t(подпись)\t',)
            run = p.add_run(f'\t',)
            run.font.underline = False
            run.font.italic = False
            run = p.add_run(f'\t(дата)\t',)
            tabs = p.paragraph_format.tab_stops
            tabs.add_tab_stop(self.cm(4.5)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(9)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(10)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(12)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(14)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(self.cm(15)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(
                self.cm(16.75)).alignment = self.wd_tb_align.CENTER
            tabs.add_tab_stop(
                self.cm(18.5)).alignment = self.wd_tb_align.CENTER
        fill_podpis2(fio, 'SK', general)

     # Сохранение документа
        print(
            "Сохранение АВК:   " + "VK_" + general['file_name']
            + "_" + self.dtime.strftime('%m%d-%H-%M-%S') + ".docx\n")
        self.document.save(
            path + "VK_" + general['file_name']
            + "_" + self.dtime.strftime('%m%d-%H-%M-%S') + ".docx")

    def insert_1(self, txt_title, txt_description, text, end):
        """ Вставка текста \n
        txt_title - текст заголовка \n
        txt_description - текст описания \n
        text - добавляемый текст\n
        """
        def spisok(var):
            if isinstance(var, str):
                return var
            else:
                if isinstance(var[0], str):
                    return var[0]
                else:
                    if isinstance(var[0][0], str):
                        return var[0][0]
        txt = spisok(text)
        self.p = self.document.add_paragraph(txt_title, style='style_zag')
        self.p = self.document.add_paragraph(f"{txt}{end}", style='style_txt')
        self.p = self.document.add_paragraph(txt_description, style='style_ops')

    def get_one_search(self, query):
        """ Поиска номера акта и даты """
        text = self.db.db_query(query)
        text_i = ''
        for i in text:
            for j in i:
                text_i += j
        return text_i

    def get_fio(self, col, sel, act):
        """ Поиска номера акта и даты\n
        col - Название возвращаемого столбца\n
        sel - Код ZK, PD\n
        act - Нормер акта\n
        """
        # Номер
        num_act = f"SELECT fio FROM АОСР WHERE number=='{act}' LIMIT 1;"
        txt_act = self.db.db_query(num_act)[0][0]
        # Данные
        sel_act = f"SELECT {col} FROM ФИО WHERE id1=='{txt_act}' AND id2=='{sel}';"
        table = self.db.db_query(sel_act)
        # Сбор текста
        text_i = ''
        for i in table:
            for j in i:
                text_i += j
                text_i += "\t "
        return text_i


if __name__ in "__main__":
    fill_act = Filling_act()
    # АОСР
    path = 'files\\'
    acts = ['680АР-02']
    for act in acts:
        fill_act.creat_aosr(path, act)